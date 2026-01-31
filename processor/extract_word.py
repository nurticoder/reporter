from __future__ import annotations

import re
from dataclasses import dataclass
from datetime import date
from typing import Any

from docx import Document

MONTHS = {
    "январь": 1,
    "января": 1,
    "февраль": 2,
    "февраля": 2,
    "март": 3,
    "марта": 3,
    "апрель": 4,
    "апреля": 4,
    "май": 5,
    "мая": 5,
    "июнь": 6,
    "июня": 6,
    "июль": 7,
    "июля": 7,
    "август": 8,
    "августа": 8,
    "сентябрь": 9,
    "сентября": 9,
    "октябрь": 10,
    "октября": 10,
    "ноябрь": 11,
    "ноября": 11,
    "декабрь": 12,
    "декабря": 12,
}

REPORT_MONTH_PATTERNS = [
    re.compile(r"(\d{4})\s*[-–]?\s*жыл\w*\s+([А-Яа-яӨөҮүҢңІіЁё]+)", re.IGNORECASE),
    re.compile(r"за\s+([А-Яа-яЁё]+)\s+(\d{4})\s*г", re.IGNORECASE),
]

DATE_PATTERN = re.compile(
    r"\b(\d{1,2})[./-](\d{1,2})[./-](\d{2,4})(?:\s*г\.?|\.?ж\.?)?\b"
)
ISO_DATE_PATTERN = re.compile(r"\b(\d{4})-(\d{2})-(\d{2})\b")
KY_DATE_PATTERN_1 = re.compile(
    r"\b(\d{4})-жыл(?:дын)?\s*(\d{1,2})[-\s]*([А-Яа-яӨөҮүҢңІіЁё]+)\s*күнү\b",
    re.IGNORECASE,
)
KY_DATE_PATTERN_2 = re.compile(
    r"\b(\d{4})-жыл(?:дын)?\s*([А-Яа-яӨөҮүҢңІіЁё]+)\s*(\d{1,2})\s*күнү\b",
    re.IGNORECASE,
)
CASE_ID_PATTERN = re.compile(
    r"(?:ЕРП|КЖБР)\s*№?\s*([\d\-/]+)", re.IGNORECASE
)
ALT_CASE_ID_PATTERN = re.compile(r"\b\d{2}-\d{3}-\d{4}-\d{6}\b")
FALLBACK_CASE_ID_PATTERN = re.compile(r"\b\d{6,}\b")
ARTICLE_PATTERN = re.compile(
    r"(?:ст\.?|статья)\s*([0-9]{1,3})(?:\s*[-–]\s*([0-9]+))?(?:\s*ч\.?\s*([0-9]+))?",
    re.IGNORECASE,
)
ARTICLE_UK_PATTERN = re.compile(
    r"\b([0-9]{1,3})(?:\s*[-–]\s*([0-9]+))?(?:\s*ч\.?\s*([0-9]+))?\s*(?:УК|КК)\b",
    re.IGNORECASE,
)
ARTICLE_KG_PATTERN = re.compile(
    r"\b([0-9]{1,3})\s*[-–]?\s*бер(?:\s*([0-9]+)\s*[-–]?\s*б\.)?\b",
    re.IGNORECASE,
)
ARTICLE_GENERIC_PATTERN = re.compile(
    r"(?:ст\.?|бер)\s*([0-9]{1,3})(?:\s*[-–]\s*([0-9]{1,3}))?(?:\s*[-–]?\s*([0-9]+)\s*[-–]?\s*б\.)?",
    re.IGNORECASE,
)

WOMEN_TAGS = [
    "аялга карата",
    "аялдар",
    "аялга",
    "аял",
]

MINOR_TAGS = [
    "жашы жетпеген",
    "жашы жетпегендер",
    "жаш өспүрүм",
    "өспүрүм",
    "балдар",
    "балдарга",
]

STOP_WORDS = [
    "токт",
    "токтог",
    "токтот",
    "токтоду",
    "токтогон",
    "кыскар",
    "кыскарт",
    "кыскарган",
    "кыскартылган",
]


@dataclass
class SourcePointer:
    table_index: int
    row_index: int
    col_index: int
    text: str


def normalize_text(value: str) -> str:
    cleaned = value.replace("\u00a0", " ").replace("–", "-").replace("—", "-")
    cleaned = re.sub(r"\s+", " ", cleaned)
    return cleaned.strip()


def normalize_lower(value: str) -> str:
    return normalize_text(value).lower()


def parse_report_month(texts: list[str]) -> dict | None:
    for text in texts:
        for pattern in REPORT_MONTH_PATTERNS:
            match = pattern.search(text)
            if not match:
                continue
            if pattern is REPORT_MONTH_PATTERNS[0]:
                year = int(match.group(1))
                month_name = match.group(2).lower()
            else:
                month_name = match.group(1).lower()
                year = int(match.group(2))
            month = MONTHS.get(month_name)
            if month:
                return {"year": year, "month": month, "label": f"{year:04d}-{month:02d}"}
    return None


def parse_date_from_text(text: str) -> date | None:
    match = DATE_PATTERN.search(text)
    if match:
        day, month, year = match.groups()
        if len(year) == 2:
            year = f"20{year}"
        try:
            return date(int(year), int(month), int(day))
        except ValueError:
            return None

    match = ISO_DATE_PATTERN.search(text)
    if match:
        year, month, day = match.groups()
        try:
            return date(int(year), int(month), int(day))
        except ValueError:
            return None

    match = KY_DATE_PATTERN_1.search(text)
    if match:
        year, day, month_name = match.groups()
        month = MONTHS.get(month_name.lower())
        if month:
            try:
                return date(int(year), int(month), int(day))
            except ValueError:
                return None

    match = KY_DATE_PATTERN_2.search(text)
    if match:
        year, month_name, day = match.groups()
        month = MONTHS.get(month_name.lower())
        if month:
            try:
                return date(int(year), int(month), int(day))
            except ValueError:
                return None

    return None


def extract_metrics(document: Document, metric_dictionary: dict) -> tuple[dict, list[dict], list[dict]]:
    metrics: dict[str, dict[str, Any]] = {}
    issues: list[dict] = []

    def iter_metric_sources():
        for idx, paragraph in enumerate(document.paragraphs):
            raw_text = paragraph.text.strip()
            if not raw_text:
                continue
            text = normalize_text(raw_text)
            yield {
                "text": text,
                "source": f"paragraph {idx + 1}",
                "cells": None,
                "kind": "paragraph",
            }

        for t_index, table in enumerate(document.tables):
            for r_index, row in enumerate(table.rows):
                cell_texts = [normalize_text(cell.text) for cell in row.cells]
                if not any(cell_texts):
                    continue
                row_text = " | ".join(cell_texts)
                yield {
                    "text": row_text,
                    "source": f"table {t_index + 1} row {r_index + 1}",
                    "cells": cell_texts,
                    "kind": "table",
                }

    def find_number(text: str) -> str | None:
        match = re.search(r"\b(\d[\d\s]*)\b", text)
        return match.group(1) if match else None

    for source in iter_metric_sources():
        text = source["text"]
        lowered = text.lower()

        for key, entry in metric_dictionary.items():
            if source["kind"] == "table" and not entry.get("allowTable"):
                continue
            found_value = None
            for regex in entry.get("regex", []):
                match = re.search(regex, text, re.IGNORECASE)
                if match:
                    found_value = match.group(1)
                    break

            if found_value is None:
                matched_phrase = False
                matched_cells: list[int] = []
                for phrase in entry.get("patterns", []):
                    phrase_norm = normalize_lower(phrase)
                    if phrase_norm in lowered:
                        matched_phrase = True
                        break
                    if source["cells"]:
                        for cell_index, cell_text in enumerate(source["cells"]):
                            if phrase_norm in normalize_lower(cell_text):
                                matched_phrase = True
                                matched_cells.append(cell_index)
                                break
                    if matched_phrase:
                        break

                if matched_phrase:
                    if source["cells"]:
                        for cell_index, cell_text in enumerate(source["cells"]):
                            if cell_index in matched_cells:
                                continue
                            found_value = find_number(cell_text)
                            if found_value:
                                break
                        if found_value is None:
                            for cell_text in source["cells"]:
                                found_value = find_number(cell_text)
                                if found_value:
                                    break
                    if found_value is None:
                        found_value = find_number(text)

            if found_value is None:
                continue

            try:
                value = int(str(found_value).replace(" ", ""))
            except ValueError:
                issues.append(
                    {
                        "type": "error",
                        "message": f"Non-numeric value for metric {key}.",
                        "source": source["source"],
                        "suggestedFix": "Ensure the metric value is a number.",
                    }
                )
                continue
            if key in metrics:
                issues.append(
                    {
                        "type": "warning",
                        "message": f"Duplicate metric found for {key}; using latest occurrence.",
                        "source": source["source"],
                        "suggestedFix": "Verify duplicate metric blocks in the Word report.",
                    }
                )
            metrics[key] = {
                "value": value,
                "sourceSnippet": text,
                "sourcePointer": source["source"],
            }

    metrics_list = [
        {
            "key": key,
            "value": data["value"],
            "sourceSnippet": data["sourceSnippet"],
            "sourcePointer": data["sourcePointer"],
        }
        for key, data in metrics.items()
    ]

    return metrics, metrics_list, issues


def extract_cases(document: Document) -> tuple[list[dict], list[dict], list[dict], int]:
    cases: list[dict] = []
    issues: list[dict] = []
    warnings: list[dict] = []
    table_count = len(document.tables)
    duplicate_case_ids: list[str] = []
    total_case_rows = 0

    for t_index, table in enumerate(document.tables):
        table_seen_ids: set[str] = set()
        for r_index, row in enumerate(table.rows):
            cell_texts = [normalize_text(cell.text) for cell in row.cells]
            if not any(cell_texts):
                continue
            row_text = " | ".join(cell_texts)
            combined_text = " ".join(cell_texts[:6]) if len(cell_texts) >= 6 else row_text

            case_id = None
            case_col = None
            for c_index, cell_text in enumerate(cell_texts):
                match = CASE_ID_PATTERN.search(cell_text)
                if match:
                    case_id = match.group(1)
                    case_col = c_index
                    break
            if case_id is None:
                match = ALT_CASE_ID_PATTERN.search(row_text)
                if match:
                    case_id = match.group(0)
                    case_col = None

            if case_id is None:
                for c_index, cell_text in enumerate(cell_texts):
                    match = FALLBACK_CASE_ID_PATTERN.search(cell_text)
                    if match:
                        case_id = match.group(0)
                        case_col = c_index
                        break

            if case_id is None:
                continue

            total_case_rows += 1
            if case_id in table_seen_ids:
                duplicate_case_ids.append(case_id)
                continue
            table_seen_ids.add(case_id)

            registered_date = None
            date_col = None
            for c_index, cell_text in enumerate(cell_texts):
                parsed = parse_date_from_text(cell_text)
                if parsed:
                    registered_date = parsed
                    date_col = c_index
                    break
            if registered_date is None:
                parsed = parse_date_from_text(combined_text)
                if parsed:
                    registered_date = parsed

            article_match = None
            article_col = None
            article_match_kind = None
            for c_index, cell_text in enumerate(cell_texts):
                match = ARTICLE_PATTERN.search(cell_text)
                if match:
                    article_match = match
                    article_col = c_index
                    article_match_kind = "std"
                    break
            if article_match is None:
                for c_index, cell_text in enumerate(cell_texts):
                    match = ARTICLE_UK_PATTERN.search(cell_text)
                    if match:
                        article_match = match
                        article_col = c_index
                        article_match_kind = "uk"
                        break
            if article_match is None:
                for c_index, cell_text in enumerate(cell_texts):
                    match = ARTICLE_KG_PATTERN.search(cell_text)
                    if match:
                        article_match = match
                        article_col = c_index
                        article_match_kind = "kg"
                        break
            if article_match is None:
                match = ARTICLE_GENERIC_PATTERN.search(combined_text)
                if match:
                    article_match = match
                    article_col = None
                    article_match_kind = "generic"
            if article_match is None:
                match = ARTICLE_GENERIC_PATTERN.search(row_text)
                if match:
                    article_match = match
                    article_col = None
                    article_match_kind = "generic"

            if registered_date is None:
                warnings.append(
                    {
                        "type": "warning",
                        "message": f"Missing registered date for case {case_id}.",
                        "source": f"table {t_index + 1} row {r_index + 1}",
                        "suggestedFix": "Ensure each case row contains a registration date.",
                    }
                )

            if article_match is None:
                warnings.append(
                    {
                        "type": "warning",
                        "message": f"Missing article code for case {case_id}.",
                        "source": f"table {t_index + 1} row {r_index + 1}",
                        "suggestedFix": "Ensure each case row contains a 'бер' or 'ст.' reference.",
                    }
                )

            article_base = None
            article_suffix = ""
            article_display = ""
            if article_match:
                article_base = article_match.group(1)
                if article_match_kind in {"std", "uk"}:
                    if article_match.group(2):
                        article_suffix += f"-{article_match.group(2)}"
                    if article_match.group(3):
                        article_suffix += f" ч.{article_match.group(3)}"
                elif article_match_kind == "kg":
                    if article_match.group(2):
                        article_suffix += f" ч.{article_match.group(2)}"
                elif article_match_kind == "generic":
                    if article_match.group(2):
                        article_suffix += f"-{article_match.group(2)}"
                    if article_match.group(3):
                        article_suffix += f" ч.{article_match.group(3)}"
                article_display = f"ст.{article_base}{article_suffix}"

            outcome_text = ""
            outcome_col = None
            if len(cell_texts) > 6 and cell_texts[6]:
                outcome_text = cell_texts[6]
                outcome_col = 6
            else:
                for c_index, cell_text in enumerate(cell_texts):
                    lowered = cell_text.lower()
                    if "токт" in lowered or "246-" in lowered or "прокур" in lowered:
                        outcome_text = cell_text
                        outcome_col = c_index
                        break
                if not outcome_text:
                    outcome_text = row_text

            description_text = " ".join(cell_texts[2:4]) if len(cell_texts) >= 4 else row_text
            normalized_desc = normalize_lower(description_text)
            women_tag = any(tag in normalized_desc for tag in WOMEN_TAGS)
            minor_tag = any(tag in normalized_desc for tag in MINOR_TAGS)

            sources = []
            for c_index, cell_text in enumerate(cell_texts):
                sources.append(
                    {
                        "table_index": t_index + 1,
                        "row_index": r_index + 1,
                        "col_index": c_index + 1,
                        "text": cell_text,
                    }
                )

            cases.append(
                {
                    "case_id": case_id,
                    "registered_date": registered_date.isoformat()
                    if registered_date
                    else None,
                    "article_base": article_base,
                    "article_suffix": article_suffix,
                    "article_display": article_display,
                    "outcome": outcome_text,
                    "tags": {"women": women_tag, "minor": minor_tag},
                    "source_cells": sources,
                    "case_id_source": {
                        "table_index": t_index + 1,
                        "row_index": r_index + 1,
                        "col_index": (case_col + 1) if case_col is not None else None,
                        "text": row_text,
                    },
                    "date_source": {
                        "table_index": t_index + 1,
                        "row_index": r_index + 1,
                        "col_index": (date_col + 1) if date_col is not None else None,
                        "text": row_text,
                    },
                    "article_source": {
                        "table_index": t_index + 1,
                        "row_index": r_index + 1,
                        "col_index": (article_col + 1) if article_col is not None else None,
                        "text": row_text,
                    },
                    "outcome_source": {
                        "table_index": t_index + 1,
                        "row_index": r_index + 1,
                        "col_index": (outcome_col + 1) if outcome_col is not None else None,
                        "text": outcome_text,
                    },
                }
            )

    if table_count == 0:
        issues.append(
            {
                "type": "error",
                "message": "No case tables found in the Word report.",
                "source": "document",
                "suggestedFix": "Ensure the report includes the case tables.",
            }
        )

    if duplicate_case_ids:
        warnings.append(
            {
                "type": "warning",
                "code": "duplicate_case_ids",
                "message": f"Duplicate case IDs detected (deduplicated): {', '.join(sorted(set(duplicate_case_ids)))}.",
                "source": "case tables",
                "suggestedFix": "Remove duplicate case rows before retrying.",
                "duplicateCount": len(duplicate_case_ids),
                "totalCases": total_case_rows,
            }
        )

    return cases, issues, warnings, table_count


def apply_case_flags(cases: list[dict], report_month: dict | None) -> list[dict]:
    for case in cases:
        reg_date = None
        if case.get("registered_date"):
            try:
                year, month, day = case["registered_date"].split("-")
                reg_date = date(int(year), int(month), int(day))
            except ValueError:
                reg_date = None

        is_new = False
        if reg_date and report_month:
            is_new = reg_date.year == report_month["year"] and reg_date.month == report_month["month"]

        normalized_outcome = normalize_lower(case.get("outcome", ""))
        has_stop_word = any(word in normalized_outcome for word in STOP_WORDS)
        has_246 = "246-" in normalized_outcome
        stop_date_match = parse_date_from_text(case.get("outcome", ""))
        stop_in_month = False
        if stop_date_match and report_month:
            stop_in_month = (
                stop_date_match.year == report_month["year"]
                and stop_date_match.month == report_month["month"]
            )
        is_stopped = has_stop_word or (has_246 and stop_in_month)

        case["is_new"] = is_new
        case["is_stopped"] = is_stopped

    return cases


def build_article_breakdown(cases: list[dict], report_month: dict | None) -> list[dict]:
    breakdown: dict[str, dict[str, int]] = {}

    for case in cases:
        if report_month and case.get("registered_date"):
            try:
                year, month, day = case["registered_date"].split("-")
                reg_date = date(int(year), int(month), int(day))
                if reg_date.year != report_month["year"] or reg_date.month != report_month["month"]:
                    continue
            except ValueError:
                pass
        base = case.get("article_base")
        if not base:
            continue
        key = f"ст.{base}"
        if key not in breakdown:
            breakdown[key] = {
                "women_u18": 0,
                "women_ge18": 0,
                "women_total": 0,
                "stopped": 0,
                "new": 0,
                "total_cases": 0,
            }

        row = breakdown[key]
        tags = case.get("tags", {})
        women = tags.get("women", False)
        minor = tags.get("minor", False)
        if women and minor:
            row["women_u18"] += 1
        if women and not minor:
            row["women_ge18"] += 1
        if women:
            row["women_total"] += 1
        if case.get("is_stopped"):
            row["stopped"] += 1
        if case.get("is_new"):
            row["new"] += 1
        row["total_cases"] += 1

    return [
        {
            "article": key,
            **values,
        }
        for key, values in sorted(breakdown.items())
    ]


def extract_word_data(docx_path: str, metric_dictionary: dict) -> dict:
    document = Document(docx_path)
    paragraph_texts = [normalize_text(p.text) for p in document.paragraphs if p.text.strip()]

    report_month = parse_report_month(paragraph_texts)

    metrics, metrics_list, metric_issues = extract_metrics(document, metric_dictionary)
    cases, case_issues, warnings, table_count = extract_cases(document)
    cases = apply_case_flags(cases, report_month)
    article_breakdown = build_article_breakdown(cases, report_month)

    return {
        "report_month": report_month,
        "metrics": metrics,
        "metrics_list": metrics_list,
        "metric_issues": metric_issues,
        "cases": cases,
        "case_issues": case_issues,
        "warnings": warnings,
        "table_count": table_count,
        "article_breakdown": article_breakdown,
    }
