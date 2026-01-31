from __future__ import annotations

from pathlib import Path

from docx import Document
from openpyxl import Workbook


def create_sample_docx(path: Path) -> None:
    doc = Document()
    doc.add_paragraph("2025-жылдын НОЯБРЬ айындагы отчет")
    doc.add_paragraph("Өткөн айдан калган кылмыш иштер – 5")
    doc.add_paragraph("Козголгон кылмыш иштер – 10")
    doc.add_paragraph("Башка органдардан келип түшкөн өндүрүштөр – 2")
    doc.add_paragraph("өзүнчө өндүрүшкө бөлүнүп алынгандар – 1")
    doc.add_paragraph("мүчүлүштүктөрүн жоюу үчүн кайтарылганы – 0")
    doc.add_paragraph("Кылмыш иштердин бириктирилгени – 1")
    doc.add_paragraph("ПРОКУРОРГО ЖӨНӨТҮЛГӨНҮ (айыптоо актысы менен) (ст.256) – 3")
    doc.add_paragraph("ПРОКУРОРГО ЖӨНӨТҮЛГӨНҮ (айыптоо актысы менен) (ст.487) – 1")
    doc.add_paragraph("Өндүрүштөн кыскартылганы – 2")
    doc.add_paragraph("Кылмыш иштин токтотулганы – 1")
    doc.add_paragraph("Өндүрүштө калган кылмыш иштер – 12")
    doc.add_paragraph("шектүү аныкталган өндүрүш – 8")
    doc.add_paragraph("шектүү аныкталбаган өндүрүш – 4")
    doc.add_paragraph("шектүү катары кармалгандар – 6")
    doc.add_paragraph("анын ичинен аялдар – 2")
    doc.add_paragraph("анын ичинен жашы жетпегендер – 1")

    table = doc.add_table(rows=1, cols=4)
    header = table.rows[0].cells
    header[0].text = "ЕРП №"
    header[1].text = "Дата"
    header[2].text = "Статья"
    header[3].text = "Примечание"

    row1 = table.add_row().cells
    row1[0].text = "ЕРП №123456"
    row1[1].text = "15.11.2025"
    row1[2].text = "ст.154"
    row1[3].text = "аялга карата токтотулган 20.11.2025"

    row2 = table.add_row().cells
    row2[0].text = "КЖБР №789012"
    row2[1].text = "05.11.2025"
    row2[2].text = "ст.209-1"
    row2[3].text = "аялга карата жашы жетпеген"

    doc.save(path)


def create_sample_docx_missing_metric(path: Path) -> None:
    doc = Document()
    doc.add_paragraph("за ноябрь 2025 г.")
    doc.add_paragraph("Өткөн айдан калган кылмыш иштер – 5")
    doc.add_paragraph("Козголгон кылмыш иштер – 10")
    doc.add_paragraph("Башка органдардан келип түшкөн өндүрүштөр – 2")
    doc.add_paragraph("өзүнчө өндүрүшкө бөлүнүп алынгандар – 1")
    doc.add_paragraph("мүчүлүштүктөрүн жоюу үчүн кайтарылганы – 0")
    doc.add_paragraph("ПРОКУРОРГО ЖӨНӨТҮЛГӨНҮ (айыптоо актысы менен) (ст.256) – 3")
    doc.add_paragraph("ПРОКУРОРГО ЖӨНӨТҮЛГӨНҮ (айыптоо актысы менен) (ст.487) – 1")
    doc.add_paragraph("Өндүрүштөн кыскартылганы – 2")
    doc.add_paragraph("Кылмыш иштин токтотулганы – 1")
    doc.add_paragraph("Өндүрүштө калган кылмыш иштер – 12")
    doc.add_paragraph("шектүү аныкталган өндүрүш – 8")
    doc.add_paragraph("шектүү аныкталбаган өндүрүш – 4")
    doc.add_paragraph("шектүү катары кармалгандар – 6")
    doc.add_paragraph("анын ичинен аялдар – 2")
    doc.add_paragraph("анын ичинен жашы жетпегендер – 1")
    doc.save(path)


def create_sample_xlsx(path: Path) -> None:
    wb = Workbook()
    ws = wb.active
    ws.title = "Отчет 1-Е Р.2"
    ws["A1"] = "Показатель"
    ws["B1"] = "Код"
    ws["C1"] = "МВД"

    labels = [
        "т. 2.1",
        "т. 2.2",
        "т. 2.3",
        "т. 2.4",
        "т. 2.5",
        "т. 2.6",
        "т. 2.7.1",
        "т. 2.7.2",
        "т. 2.8",
        "т. 2.9",
        "т. 2.10",
        "т. 3.1",
        "т. 3.2",
        "т. 4.1",
        "т. 4.1 (аялдар)",
        "т. 4.1 (жашы жетпегендер)",
        "ст.154",
        "ст.209",
    ]

    start_row = 2
    for i, label in enumerate(labels):
        ws.cell(row=start_row + i, column=2, value=label)
        ws.cell(row=start_row + i, column=3, value=0)

    wb.save(path)
