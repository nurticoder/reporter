from __future__ import annotations

import re
from dataclasses import dataclass
from datetime import date
from typing import Any

from docx import Document

from .normalize import (
    MONTHS,
    STOP_WORDS,
    contains_any,
    normalize_article,
    normalize_text,
    parse_date_from_text,
    parse_int,
)

REPORT_MONTH_PATTERNS = [
    re.compile(r"(\d{4})\s*[-–]?\s*жыл\w*\s+([А-Яа-яӨөҮүҢңІіЁё]+)", re.IGNORECASE),
    re.compile(r"за\s+([А-Яа-яЁё]+)\s+(\d{4})\s*г", re.IGNORECASE),
]

CASE_ID_PATTERN = re.compile(r"(?:ЕРП|КЖБР)\s*№?\s*([\d\-/]+)", re.IGNORECASE)
ALT_CASE_ID_PATTERN = re.compile(r"\b\d{2}-\d{3}-\d{4}-\d{6}\b")
FALLBACK_CASE_ID_PATTERN = re.compile(r"\b\d{6,}\b")

WOMEN_TAGS = ["аялга карата", "аялдар", "аялга", "аял"]
MINOR_TAGS = ["жашы жетпеген", "жашы жетпегендер", "жаш өспүрүм", "өспүрүм", "балдар", "балдарга"]


@dataclass
class MetricHit:
    key: str
    value: int
    source: str
    snippet: str


def parse_report_month(paragraphs: list[str]) -> dict | None:
    for text in paragraphs:
        for pattern in REPORT_MONTH_PATTERNS:
            match = pattern.search(text)
            if not match:
                continue
            if pattern is REPORT_MONTH_PATTERNS[0]:
                year = int(match.group(1))
                month_name = match.group(2).lower()
            else:
                month_name = match.group(1).lower()
                year = int(match.group(2))
            month = MONTHS.get(month_name)
            if month:
                return {"year": year, "month": month, "label": f"{year:04d}-{month:02d}"}
    return None


def extract_metrics(paragraphs: list[str], metrics_config: dict) -> tuple[dict, list[dict], list[dict]]:
    metrics: dict[str, dict[str, Any]] = {}
    warnings: list[dict] = []

    for idx, raw in enumerate(paragraphs):
        text = normalize_text(raw)
        lowered = text.lower()
        for key, entry in metrics_config.items():
            found = None
            for pattern in entry.get("regex", []):
                match = re.search(pattern, text, re.IGNORECASE)
                if match:
                    found = match.group(1)
                    break
            if found is None:
                for phrase in entry.get("phrases", []):
                    if phrase.lower() in lowered:
                        found = re.search(r"(\d[\d\s]*)", text)
                        found = found.group(1) if found else None
                        break
            if found is None:
                continue
            value = parse_int(found)
            if value is None:
                continue
            if key in metrics:
                warnings.append(
                    {
                        "type": "warning",
                        "message": f"Duplicate metric found for {key}; using latest occurrence.",
                        "source": f"paragraph {idx + 1}",
                        "suggestedFix": "Verify duplicate metric blocks in the Word report.",
                        "previous": metrics[key],
                        "latestSnippet": text,
                    }
                )
            metrics[key] = {
                "value": value,
                "sourceSnippet": text,
                "sourcePointer": f"paragraph {idx + 1}",
            }

    metrics_list = [
        {
            "key": key,
            "value": data["value"],
            "sourceSnippet": data["sourceSnippet"],
            "sourcePointer": data["sourcePointer"],
        }
        for key, data in metrics.items()
    ]

    return metrics, metrics_list, warnings


def extract_cases(document: Document) -> tuple[list[dict], list[dict]]:
    cases: list[dict] = []
    warnings: list[dict] = []

    for t_index, table in enumerate(document.tables):
        table_seen: dict[str, int] = {}
        for r_index, row in enumerate(table.rows):
            cell_texts = [normalize_text(cell.text) for cell in row.cells]
            if not any(cell_texts):
                continue
            row_text = " | ".join(cell_texts)
            combined_text = " ".join(cell_texts[:6]) if len(cell_texts) >= 6 else row_text

            case_id = None
            for cell_text in cell_texts:
                match = CASE_ID_PATTERN.search(cell_text)
                if match:
                    case_id = match.group(1)
                    break
            if case_id is None:
                match = ALT_CASE_ID_PATTERN.search(row_text)
                if match:
                    case_id = match.group(0)
            if case_id is None:
                match = FALLBACK_CASE_ID_PATTERN.search(row_text)
                case_id = match.group(0) if match else None
            if case_id is None:
                continue

            if case_id in table_seen:
                warnings.append(
                    {
                        "type": "warning",
                        "message": f"Duplicate case ID {case_id} in table {t_index + 1}; keeping last occurrence.",
                        "source": f"table {t_index + 1}",
                        "suggestedFix": "Remove duplicate case rows before retrying.",
                    }
                )
                # remove previous occurrence
                cases = [c for c in cases if c.get("case_id") != case_id or c.get("table_index") != t_index + 1]
            table_seen[case_id] = r_index

            registered_date = parse_date_from_text(combined_text)
            article = normalize_article(row_text)

            outcome = ""
            if len(cell_texts) > 6 and cell_texts[6]:
                outcome = cell_texts[6]
            else:
                outcome = row_text

            tags_text = " ".join(cell_texts[2:4]) if len(cell_texts) >= 4 else row_text
            women_tag = contains_any(tags_text, WOMEN_TAGS)
            minor_tag = contains_any(tags_text, MINOR_TAGS)

            cases.append(
                {
                    "case_id": case_id,
                    "registered_date": registered_date.isoformat() if registered_date else None,
                    "article": article,
                    "outcome": outcome,
                    "tags": {"women": women_tag, "minor": minor_tag},
                    "table_index": t_index + 1,
                    "row_index": r_index + 1,
                }
            )

    return cases, warnings


def apply_case_flags(cases: list[dict], report_month: dict | None) -> list[dict]:
    for case in cases:
        reg_date = None
        if case.get("registered_date"):
            try:
                year, month, day = case["registered_date"].split("-")
                reg_date = date(int(year), int(month), int(day))
            except ValueError:
                reg_date = None

        is_new = False
        if reg_date and report_month:
            is_new = reg_date.year == report_month["year"] and reg_date.month == report_month["month"]

        normalized_outcome = normalize_text(case.get("outcome", "")).lower()
        has_stop_word = any(word in normalized_outcome for word in STOP_WORDS)
        is_stopped = has_stop_word

        case["is_new"] = is_new
        case["is_stopped"] = is_stopped
    return cases


def build_article_breakdown(cases: list[dict], report_month: dict | None) -> list[dict]:
    breakdown: dict[str, dict[str, int]] = {}

    for case in cases:
        if report_month and case.get("registered_date"):
            try:
                year, month, day = case["registered_date"].split("-")
                if int(year) != report_month["year"] or int(month) != report_month["month"]:
                    continue
            except ValueError:
                pass
        article = case.get("article")
        if not article:
            continue
        if article not in breakdown:
            breakdown[article] = {
                "women_u18": 0,
                "women_ge18": 0,
                "women_total": 0,
                "stopped": 0,
                "new": 0,
                "total_cases": 0,
            }
        row = breakdown[article]
        women = case.get("tags", {}).get("women", False)
        minor = case.get("tags", {}).get("minor", False)
        if women and minor:
            row["women_u18"] += 1
        if women and not minor:
            row["women_ge18"] += 1
        if case.get("is_stopped"):
            row["stopped"] += 1
        if case.get("is_new"):
            row["new"] += 1
        row["total_cases"] += 1

    for row in breakdown.values():
        row["women_total"] = row["women_u18"] + row["women_ge18"]

    return [
        {
            "article": key,
            **values,
        }
        for key, values in sorted(breakdown.items())
    ]


def extract_docx(docx_path: str, metrics_config: dict) -> dict:
    document = Document(docx_path)
    paragraphs = [normalize_text(p.text) for p in document.paragraphs if p.text.strip()]

    report_month = parse_report_month(paragraphs)

    metrics, metrics_list, metric_warnings = extract_metrics(paragraphs, metrics_config)
    cases, case_warnings = extract_cases(document)
    cases = apply_case_flags(cases, report_month)
    article_breakdown = build_article_breakdown(cases, report_month)

    return {
        "report_month": report_month,
        "metrics": metrics,
        "metrics_list": metrics_list,
        "metric_warnings": metric_warnings,
        "cases": cases,
        "case_warnings": case_warnings,
        "article_breakdown": article_breakdown,
    }
